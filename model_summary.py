#!/usr/bin/env python3.12
"""
Print a compact summary of both models showing each layer,
output shape, and parameter count.
Run: python3.12 model_summary.py
"""

import os
import torch
import torch.nn as nn
import torch.nn.functional as F
import torchvision.models as tv_models
from torchinfo import summary


class ConvBlock(nn.Module):
    def __init__(self, in_ch, out_ch):
        super().__init__()
        self.block = nn.Sequential(
            nn.Conv2d(in_ch, out_ch, 3, padding=1, bias=False),
            nn.BatchNorm2d(out_ch), nn.ReLU(inplace=True),
            nn.Conv2d(out_ch, out_ch, 3, padding=1, bias=False),
            nn.BatchNorm2d(out_ch), nn.ReLU(inplace=True),
        )
    def forward(self, x): return self.block(x)


class Up(nn.Module):
    def __init__(self, in_ch, skip_ch, out_ch):
        super().__init__()
        self.up   = nn.ConvTranspose2d(in_ch, in_ch // 2, 2, stride=2)
        self.conv = ConvBlock(in_ch // 2 + skip_ch, out_ch)
    def forward(self, x, skip):
        x = self.up(x)
        if x.shape[-2:] != skip.shape[-2:]:
            x = F.interpolate(x, size=skip.shape[-2:], mode='bilinear', align_corners=False)
        return self.conv(torch.cat([x, skip], dim=1))


class FingerSegUNet(nn.Module):
    def __init__(self, num_classes=4):
        super().__init__()
        bb = tv_models.resnet18(weights=None)
        self.enc0 = nn.Sequential(
            nn.Conv2d(1, 64, kernel_size=7, stride=2, padding=3, bias=False),
            bb.bn1, nn.ReLU(inplace=True),
        )
        self.pool = bb.maxpool
        self.enc1 = bb.layer1
        self.enc2 = bb.layer2
        self.enc3 = bb.layer3
        self.enc4 = bb.layer4
        self.up4 = Up(512, 256, 256)
        self.up3 = Up(256, 128, 128)
        self.up2 = Up(128,  64,  64)
        self.up1 = Up(64,   64,  32)
        self.up0 = nn.Sequential(
            nn.ConvTranspose2d(32, 32, kernel_size=2, stride=2),
            ConvBlock(32, 32),
        )
        self.head = nn.Conv2d(32, num_classes, 1)

    def forward(self, x):
        e0 = self.enc0(x)
        e1 = self.enc1(self.pool(e0))
        e2 = self.enc2(e1)
        e3 = self.enc3(e2)
        b  = self.enc4(e3)
        d  = self.up4(b,  e3)
        d  = self.up3(d,  e2)
        d  = self.up2(d,  e1)
        d  = self.up1(d,  e0)
        d  = self.up0(d)
        return self.head(d)


class FingerKpModel(nn.Module):
    def __init__(self, num_kp=6):
        super().__init__()
        self.enc1   = ConvBlock(1,   32)
        self.pool1  = nn.MaxPool2d(2)
        self.enc2   = ConvBlock(32,  64)
        self.pool2  = nn.MaxPool2d(2)
        self.enc3   = ConvBlock(64,  128)
        self.pool3  = nn.MaxPool2d(2)
        self.bottle = ConvBlock(128, 256)
        self.up3    = nn.ConvTranspose2d(256, 128, 2, stride=2)
        self.dec3   = ConvBlock(256, 128)
        self.up2    = nn.ConvTranspose2d(128, 64, 2, stride=2)
        self.dec2   = ConvBlock(128, 64)
        self.up1    = nn.ConvTranspose2d(64, 32, 2, stride=2)
        self.dec1   = ConvBlock(64, 32)
        self.head   = nn.Conv2d(32, num_kp, 1)

    @staticmethod
    def _up(layer, x, ref):
        x = layer(x)
        if x.shape[-2:] != ref.shape[-2:]:
            x = F.interpolate(x, size=ref.shape[-2:], mode='bilinear', align_corners=False)
        return x

    def forward(self, x):
        e1 = self.enc1(x)
        e2 = self.enc2(self.pool1(e1))
        e3 = self.enc3(self.pool2(e2))
        b  = self.bottle(self.pool3(e3))
        d3 = self.dec3(torch.cat([self._up(self.up3, b,  e3), e3], dim=1))
        d2 = self.dec2(torch.cat([self._up(self.up2, d3, e2), e2], dim=1))
        d1 = self.dec1(torch.cat([self._up(self.up1, d2, e1), e1], dim=1))
        return self.head(d1)


if __name__ == '__main__':
    import sys
    import io
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    os.makedirs('trained_model', exist_ok=True)
    doc = Document()
    doc.add_heading('Model Architecture Summaries', 0)

    for model, name, input_size in [
        (FingerSegUNet(4), "SEG MODEL — ResNet18-UNet  (input: 1×256×256)", (1, 1, 256, 256)),
        (FingerKpModel(6), "KP MODEL  — U-Net          (input: 1×256×256)", (1, 1, 256, 256)),
    ]:
        doc.add_heading(name, level=1)

        # Capture summary output as string
        buf = io.StringIO()
        old_stdout = sys.stdout
        sys.stdout = buf
        summary(model, input_size=input_size,
                col_names=["output_size", "num_params"], depth=3)
        sys.stdout = old_stdout
        text = buf.getvalue()

        # Add as monospaced paragraph to preserve formatting
        para = doc.add_paragraph()
        run  = para.add_run(text)
        run.font.name = 'Courier New'
        run.font.size = Pt(7)
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        doc.add_paragraph()

    output_path = 'trained_model/model_summaries.docx'
    doc.save(output_path)
    print(f"Saved: {output_path}")
