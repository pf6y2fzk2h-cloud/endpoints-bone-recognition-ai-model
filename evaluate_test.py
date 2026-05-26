#!/usr/bin/env python3.12
"""
Evaluate model_11_6kp on held-out test images.
===============================================
Reads images from ../test_images/ and ground-truth from
../test_images/labels.json, then produces a Word document
with per-image tables and visualisations.

Run:
    python3.12 evaluate_test.py
"""

import sys; sys.stdout.reconfigure(line_buffering=True)
import torch
import torch.nn as nn
import torch.nn.functional as F
import torchvision.models as tv_models
from PIL import Image
import numpy as np
import cv2
import os
import json
from pathlib import Path
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches

# =============================================================================
# CONFIG
# =============================================================================

CONFIG = {
    'annotations_file': '../test_images/labels.json',
    'images_dir':       '../test_images',
    'seg_model_path':   'trained_model/best_seg_model.pth',
    'kp_model_path':    'trained_model/best_kp_model.pth',
    'output_folder':    'evaluation_test',
    'image_size':       256,
    'crop_size':        256,
    'seg_threshold':    0.2,
    'kp_threshold':     0.0,
    'pad_fraction':     0.10,
    'px_to_mm':         0.265,   # 96 DPI — 1 px = 0.265 mm
}

FINGER_LABELS = ['L2', 'L4', 'R2', 'R4']
FINGER_KP_NAMES = {
    'L2': ['L21-','L21_','L22-','L22_','L23-','L23_'],
    'L4': ['L41-','L41_','L42-','L42_','L43-','L43_'],
    'R2': ['R21-','R21_','R22-','R22_','R23-','R23_'],
    'R4': ['R41-','R41_','R42-','R42_','R43-','R43_'],
}
ALL_KP_LABELS = (FINGER_KP_NAMES['L2'] + FINGER_KP_NAMES['L4'] +
                 FINGER_KP_NAMES['R2'] + FINGER_KP_NAMES['R4'])

FINGER_COLORS = {
    'L2': (255,   0, 255),
    'L4': (255, 165,   0),
    'R2': (  0, 255, 255),
    'R4': (  0, 255,   0),
}

# =============================================================================
# MODELS
# =============================================================================

class ConvBlock(nn.Module):
    def __init__(self, in_ch, out_ch):
        super().__init__()
        self.block = nn.Sequential(
            nn.Conv2d(in_ch, out_ch, 3, padding=1, bias=False),
            nn.BatchNorm2d(out_ch), nn.ReLU(inplace=True),
            nn.Conv2d(out_ch, out_ch, 3, padding=1, bias=False),
            nn.BatchNorm2d(out_ch), nn.ReLU(inplace=True),
        )
    def forward(self, x): return self.block(x)


class Up(nn.Module):
    def __init__(self, in_ch, skip_ch, out_ch):
        super().__init__()
        self.up   = nn.ConvTranspose2d(in_ch, in_ch // 2, 2, stride=2)
        self.conv = ConvBlock(in_ch // 2 + skip_ch, out_ch)
    def forward(self, x, skip):
        x = self.up(x)
        if x.shape[-2:] != skip.shape[-2:]:
            x = F.interpolate(x, size=skip.shape[-2:], mode='bilinear', align_corners=False)
        return self.conv(torch.cat([x, skip], dim=1))


class FingerSegUNet(nn.Module):
    def __init__(self, num_classes=4, pretrained=False):
        super().__init__()
        try:
            bb = tv_models.resnet18(weights='IMAGENET1K_V1' if pretrained else None)
        except TypeError:
            bb = tv_models.resnet18(pretrained=pretrained)
        self.enc0 = nn.Sequential(
            nn.Conv2d(1, 64, kernel_size=7, stride=2, padding=3, bias=False),
            bb.bn1, nn.ReLU(inplace=True),
        )
        if pretrained:
            with torch.no_grad():
                self.enc0[0].weight.data = bb.conv1.weight.data.mean(dim=1, keepdim=True)
        self.pool = bb.maxpool
        self.enc1 = bb.layer1; self.enc2 = bb.layer2
        self.enc3 = bb.layer3; self.enc4 = bb.layer4
        self.up4 = Up(512, 256, 256); self.up3 = Up(256, 128, 128)
        self.up2 = Up(128,  64,  64); self.up1 = Up(64,   64,  32)
        self.up0 = nn.Sequential(
            nn.ConvTranspose2d(32, 32, kernel_size=2, stride=2), ConvBlock(32, 32))
        self.head = nn.Conv2d(32, num_classes, 1)

    def forward(self, x):
        e0 = self.enc0(x); e1 = self.enc1(self.pool(e0))
        e2 = self.enc2(e1); e3 = self.enc3(e2); b = self.enc4(e3)
        d = self.up4(b, e3); d = self.up3(d, e2); d = self.up2(d, e1)
        d = self.up1(d, e0); d = self.up0(d)
        return self.head(d)


class FingerKpModel(nn.Module):
    def __init__(self, num_kp=6):
        super().__init__()
        self.enc1   = ConvBlock(1,   32); self.pool1 = nn.MaxPool2d(2)
        self.enc2   = ConvBlock(32,  64); self.pool2 = nn.MaxPool2d(2)
        self.enc3   = ConvBlock(64, 128); self.pool3 = nn.MaxPool2d(2)
        self.bottle = ConvBlock(128, 256)
        self.up3    = nn.ConvTranspose2d(256, 128, 2, stride=2)
        self.dec3   = ConvBlock(256, 128)
        self.up2    = nn.ConvTranspose2d(128, 64, 2, stride=2)
        self.dec2   = ConvBlock(128, 64)
        self.up1    = nn.ConvTranspose2d(64, 32, 2, stride=2)
        self.dec1   = ConvBlock(64, 32)
        self.head   = nn.Conv2d(32, num_kp, 1)

    @staticmethod
    def _up(layer, x, ref):
        x = layer(x)
        if x.shape[-2:] != ref.shape[-2:]:
            x = F.interpolate(x, size=ref.shape[-2:], mode='bilinear', align_corners=False)
        return x

    def forward(self, x):
        e1 = self.enc1(x); e2 = self.enc2(self.pool1(e1))
        e3 = self.enc3(self.pool2(e2)); b = self.bottle(self.pool3(e3))
        d3 = self.dec3(torch.cat([self._up(self.up3, b,  e3), e3], dim=1))
        d2 = self.dec2(torch.cat([self._up(self.up2, d3, e2), e2], dim=1))
        d1 = self.dec1(torch.cat([self._up(self.up1, d2, e1), e1], dim=1))
        return self.head(d1)

# =============================================================================
# PREPROCESSING + CROP UTILS
# =============================================================================

def apply_clahe(arr):
    return cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8,8)).apply(arr)

def preprocess_full(pil_gray, size):
    arr = apply_clahe(np.array(pil_gray, dtype=np.uint8))
    arr = cv2.resize(arr, (size, size))
    return torch.from_numpy(arr.astype(np.float32)/255.0).unsqueeze(0).unsqueeze(0)

def preprocess_crop(crop_arr, size, mask=None):
    arr = apply_clahe(crop_arr)
    arr = cv2.resize(arr, (size, size))
    if mask is not None:
        m     = cv2.resize(mask, (size, size))
        kern  = np.ones((9,9), np.uint8)
        m_dil = cv2.dilate((m > 0.15).astype(np.uint8), kern, iterations=3).astype(np.float32)
        arr = np.clip(arr.astype(np.float32) * m_dil, 0, 255).astype(np.uint8)
    return torch.from_numpy(arr.astype(np.float32)/255.0).unsqueeze(0).unsqueeze(0)

def mask_to_bbox(mask_np, threshold=0.2):
    binary = (mask_np > threshold).astype(np.uint8)
    if binary.sum() == 0: return None
    _, _, stats, _ = cv2.connectedComponentsWithStats(binary)
    if len(stats) <= 1: return None
    largest = 1 + np.argmax(stats[1:, cv2.CC_STAT_AREA])
    x1 = int(stats[largest, cv2.CC_STAT_LEFT])
    y1 = int(stats[largest, cv2.CC_STAT_TOP])
    x2 = x1 + int(stats[largest, cv2.CC_STAT_WIDTH])
    y2 = y1 + int(stats[largest, cv2.CC_STAT_HEIGHT])
    return (x1, y1, x2, y2)

def fix_finger_assignment(seg_probs, threshold=0.2):
    def cx(ch):
        pts = np.where(seg_probs[ch] > threshold)
        return float(pts[1].mean()) if len(pts[1]) > 0 else None
    l2_x, l4_x = cx(0), cx(1)
    if l2_x is not None and l4_x is not None and l2_x < l4_x:
        seg_probs[[0,1]] = seg_probs[[1,0]]
    r2_x, r4_x = cx(2), cx(3)
    if r2_x is not None and r4_x is not None and r2_x > r4_x:
        seg_probs[[2,3]] = seg_probs[[3,2]]
    return seg_probs

def make_square_crop(arr, bbox_model, model_size, orig_size, pad_frac=0.10):
    orig_w, orig_h = orig_size
    sx = orig_w / model_size; sy = orig_h / model_size
    x1=bbox_model[0]*sx; y1=bbox_model[1]*sy
    x2=bbox_model[2]*sx; y2=bbox_model[3]*sy
    pw=(x2-x1)*pad_frac; ph=(y2-y1)*pad_frac
    x1-=pw; y1-=ph; x2+=pw; y2+=ph
    cx=(x1+x2)/2; cy=(y1+y2)/2
    side=max(x2-x1, y2-y1)
    x1=cx-side/2; x2=cx+side/2; y1=cy-side/2; y2=cy+side/2
    cx1=int(max(0,x1)); cy1=int(max(0,y1))
    cx2=int(min(orig_w,x2)); cy2=int(min(orig_h,y2))
    if cx2<=cx1 or cy2<=cy1: return None, None
    return arr[cy1:cy2, cx1:cx2], (cx1,cy1,cx2,cy2)

# =============================================================================
# KEYPOINT EXTRACTION
# =============================================================================

def predict_crop_tta(model, crop_t, device):
    with torch.no_grad():
        hm_orig = model(crop_t.to(device))[0]
        hm_flip = model(crop_t.flip(-1).to(device))[0]
    return (hm_orig + hm_flip.flip(-1)) / 2

def extract_keypoints(heatmap_tensor, threshold=0.0):
    num_kp = heatmap_tensor.shape[0]
    hms  = torch.sigmoid(heatmap_tensor).cpu().numpy()
    kps  = np.full((num_kp, 2), np.nan, dtype=np.float32)
    conf = np.zeros(num_kp, dtype=np.float32)
    for j in range(num_kp):
        hm = hms[j]
        mv = hm.max()
        if mv < threshold:
            continue
        y, x = np.unravel_index(hm.argmax(), hm.shape)
        kps[j]  = [x / hm.shape[1], y / hm.shape[0]]
        conf[j] = mv
    return kps, conf

def crop_to_image_coords(kps_crop, crop_box):
    cx1,cy1,cx2,cy2 = crop_box
    cw=cx2-cx1; ch=cy2-cy1
    out = np.full_like(kps_crop, np.nan)
    for j in range(len(kps_crop)):
        if not np.isnan(kps_crop[j,0]):
            out[j,0] = kps_crop[j,0]*cw + cx1
            out[j,1] = kps_crop[j,1]*ch + cy1
    return out

# =============================================================================
# VISUALISATION
# =============================================================================

def _save_vis(orig_image, gt_kps, pred_kps, base_path):
    img_arr = np.array(orig_image.convert('RGB'))
    patches = [mpatches.Patch(color=np.array(v)/255.0, label=k)
               for k, v in FINGER_COLORS.items()]

    gt_path   = base_path.replace('.png', '_gt.png')
    pred_path = base_path.replace('.png', '_pred.png')

    for path, kps, title in [
        (gt_path,   gt_kps,   'Ground Truth Keypoints'),
        (pred_path, pred_kps, 'Predicted Keypoints'),
    ]:
        fig, ax = plt.subplots(1, 1, figsize=(10, 14))
        ax.imshow(img_arr, cmap='gray')
        ax.set_title(title, fontsize=13, fontweight='bold')

        for i, lbl in enumerate(ALL_KP_LABELS):
            if np.isnan(kps[i, 0]):
                continue
            fl    = lbl[:2]
            color = np.array(FINGER_COLORS[fl]) / 255.0
            ax.plot(kps[i,0], kps[i,1], 'x',
                    color=color, markersize=7, markeredgewidth=1.5, alpha=0.9)
            ax.text(kps[i,0]+6, kps[i,1]-6, lbl, color='yellow', fontsize=5,
                    fontweight='bold',
                    bbox=dict(boxstyle='round,pad=0.1', facecolor='black', alpha=0.5))

        ax.legend(handles=patches, loc='lower right', fontsize=8)
        ax.axis('off')
        fig.tight_layout()
        fig.savefig(path, dpi=150, bbox_inches='tight')
        plt.close(fig)

    return gt_path, pred_path

# =============================================================================
# MAIN
# =============================================================================

def main():
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor

    os.makedirs(CONFIG['output_folder'], exist_ok=True)
    device = torch.device('cuda' if torch.cuda.is_available() else 'cpu')

    for p in [CONFIG['seg_model_path'], CONFIG['kp_model_path']]:
        if not os.path.exists(p):
            print(f"ERROR: {p} not found. Train the models first."); return

    seg_ck = torch.load(CONFIG['seg_model_path'], map_location=device)
    seg_model = FingerSegUNet(4).to(device)
    seg_model.load_state_dict(seg_ck['model_state_dict']); seg_model.eval()

    kp_ck = torch.load(CONFIG['kp_model_path'], map_location=device)
    kp_model = FingerKpModel(6).to(device)
    kp_model.load_state_dict(kp_ck['model_state_dict']); kp_model.eval()
    print("Models loaded.\n")

    with open(CONFIG['annotations_file']) as f:
        all_data = json.load(f)

    # Match labels.json entries (may reference .jpg) to actual .png files
    selected = []
    for entry in all_data:
        stem = Path(entry['name']).stem          # "227" from "227.jpg" or "227.png"
        png_path = os.path.join(CONFIG['images_dir'], stem + '.png')
        if os.path.exists(png_path):
            selected.append((entry, stem + '.png'))

    if not selected:
        print("No test images matched in labels.json."); return

    print(f"Evaluating on {len(selected)} test images (unseen during training).\n")

    all_errors = []
    per_label_errors = {lbl: [] for lbl in ALL_KP_LABELS}

    doc = Document()
    doc.add_heading('Test Set Evaluation — model_11_6kp', 0)
    doc.add_paragraph(
        f"Model: model_11_6kp  |  "
        f"Test images: {len(selected)}  |  "
        f"Crop size: {CONFIG['crop_size']}px  |  "
        f"Keypoints: 6 per finger (24 total, bones 1–3)"
    )

    for entry, img_filename in selected:
        img_path = os.path.join(CONFIG['images_dir'], img_filename)
        print(f"Processing {img_filename}...")

        orig_image = Image.open(img_path).convert('L')
        orig_arr   = np.array(orig_image, dtype=np.uint8)
        orig_size  = orig_image.size

        # Filter to only the 24 keypoints model_11 predicts (bones 1-3)
        gt_points = {p['label']: (p['x'], p['y'])
                     for p in entry.get('points', [])
                     if p['label'] in ALL_KP_LABELS}

        seg_input = preprocess_full(orig_image, CONFIG['image_size']).to(device)
        with torch.no_grad():
            seg_probs = torch.sigmoid(seg_model(seg_input))[0].cpu().numpy()
        seg_probs = fix_finger_assignment(seg_probs, CONFIG['seg_threshold'])

        finger_masks_up = np.stack([
            cv2.resize(seg_probs[c], orig_size, interpolation=cv2.INTER_LINEAR)
            for c in range(4)], axis=0)

        pred_kps = np.full((24, 2), np.nan, dtype=np.float32)
        gt_kps   = np.full((24, 2), np.nan, dtype=np.float32)

        for ch, fl in enumerate(FINGER_LABELS):
            kp_labels = FINGER_KP_NAMES[fl]
            base = ch * 6
            sl   = slice(base, base + 6)

            bbox_model = mask_to_bbox(seg_probs[ch], CONFIG['seg_threshold'])
            if bbox_model is None: continue

            crop_arr, crop_box = make_square_crop(
                orig_arr, bbox_model, CONFIG['image_size'], orig_size,
                CONFIG['pad_fraction'])
            if crop_arr is None: continue

            cx1,cy1,cx2,cy2 = crop_box
            finger_mask = finger_masks_up[ch, cy1:cy2, cx1:cx2]

            crop_t    = preprocess_crop(crop_arr, CONFIG['crop_size'],
                                        mask=finger_mask).to(device)
            kp_logits = predict_crop_tta(kp_model, crop_t, device)
            coords_01, _ = extract_keypoints(kp_logits, CONFIG['kp_threshold'])

            kps_img = crop_to_image_coords(coords_01, crop_box)
            pred_kps[sl] = kps_img

            for j, lbl in enumerate(kp_labels):
                if lbl in gt_points:
                    gt_kps[base + j] = gt_points[lbl]

        # Collect errors
        img_errors = []
        img_rows   = []
        for i, lbl in enumerate(ALL_KP_LABELS):
            if np.isnan(gt_kps[i,0]) or np.isnan(pred_kps[i,0]):
                continue
            err = float(np.hypot(pred_kps[i,0]-gt_kps[i,0],
                                 pred_kps[i,1]-gt_kps[i,1]))
            img_errors.append(err)
            all_errors.append(err)
            per_label_errors[lbl].append(err)
            img_rows.append((lbl, gt_kps[i], pred_kps[i], err))

        img_mean = np.mean(img_errors) if img_errors else float('nan')
        print(f"  Mean error: {img_mean:.2f}px  ({len(img_errors)} keypoints)")

        doc.add_heading(img_filename, level=1)
        doc.add_paragraph(
            f"Mean error: {img_mean:.2f} px  |  Keypoints compared: {len(img_errors)}"
        )

        tbl = doc.add_table(rows=1, cols=6)
        tbl.style = 'Table Grid'
        for cell, text in zip(tbl.rows[0].cells,
                              ['Keypoint','GT x','GT y','Pred x','Pred y','Error (px)']):
            cell.text = text
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(9)

        for lbl, gt, pred, err in img_rows:
            row = tbl.add_row().cells
            row[0].text = lbl
            row[1].text = f'{gt[0]:.1f}'
            row[2].text = f'{gt[1]:.1f}'
            row[3].text = f'{pred[0]:.1f}'
            row[4].text = f'{pred[1]:.1f}'
            row[5].text = f'{err:.2f}'
            for cell in row:
                cell.paragraphs[0].runs[0].font.size = Pt(9)
            run = row[5].paragraphs[0].runs[0]
            if err < 10:
                run.font.color.rgb = RGBColor(0, 150, 0)
            elif err < 25:
                run.font.color.rgb = RGBColor(200, 100, 0)
            else:
                run.font.color.rgb = RGBColor(200, 0, 0)

        doc.add_paragraph()

        vis_base = os.path.join(CONFIG['output_folder'],
                                f'{Path(img_filename).stem}_eval.png')
        gt_path, pred_path = _save_vis(orig_image, gt_kps, pred_kps, vis_base)
        doc.add_paragraph('Ground Truth:').runs[0].bold = True
        doc.add_picture(gt_path, width=Inches(5.5))
        doc.add_paragraph('Predicted:').runs[0].bold = True
        doc.add_picture(pred_path, width=Inches(5.5))
        doc.add_paragraph()

    # Overall summary
    doc.add_heading('Overall Summary', level=1)
    if all_errors:
        mm = CONFIG['px_to_mm']
        for line in [
            f"Total keypoints compared : {len(all_errors)}",
            f"Images evaluated         : {len(selected)}",
            f"Mean error               : {np.mean(all_errors):.2f} px  "
            f"({np.mean(all_errors)*mm:.2f} mm)",
            f"Median error             : {np.median(all_errors):.2f} px  "
            f"({np.median(all_errors)*mm:.2f} mm)",
            f"Std deviation            : {np.std(all_errors):.2f} px",
            f"Min / Max                : {np.min(all_errors):.2f} / "
            f"{np.max(all_errors):.2f} px",
        ]:
            doc.add_paragraph(line)

        doc.add_heading('Per-Keypoint Mean Error', level=2)
        tbl2 = doc.add_table(rows=1, cols=3)
        tbl2.style = 'Table Grid'
        for cell, text in zip(tbl2.rows[0].cells,
                               ['Keypoint', 'Mean error (px)', 'N images']):
            cell.text = text
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(9)
        for lbl, errs in per_label_errors.items():
            if errs:
                row = tbl2.add_row().cells
                row[0].text = lbl
                row[1].text = f'{np.mean(errs):.2f}'
                row[2].text = str(len(errs))
                for cell in row:
                    cell.paragraphs[0].runs[0].font.size = Pt(9)

    docx_path = os.path.join(CONFIG['output_folder'], 'test_evaluation_results.docx')
    doc.save(docx_path)

    print("\n" + "=" * 50)
    if all_errors:
        mm = CONFIG['px_to_mm']
        print(f"TEST SET RESULTS  ({len(selected)} unseen images, {len(all_errors)} keypoints)")
        print(f"  Mean error   : {np.mean(all_errors):.2f} px  ({np.mean(all_errors)*mm:.2f} mm)")
        print(f"  Median error : {np.median(all_errors):.2f} px")
        print(f"  Std dev      : {np.std(all_errors):.2f} px")
        print(f"  Min / Max    : {np.min(all_errors):.2f} / {np.max(all_errors):.2f} px")
        print(f"\n{'Keypoint':<10}  {'Mean (px)':>10}  {'Median (px)':>12}  {'N':>4}")
        print("-" * 42)
        for lbl, errs in per_label_errors.items():
            if errs:
                print(f"  {lbl:<8}  {np.mean(errs):>10.2f}  {np.median(errs):>12.2f}  {len(errs):>4}")
    print(f"\nSaved: {docx_path}")


if __name__ == '__main__':
    try:
        main()
    except Exception:
        import traceback; traceback.print_exc()
